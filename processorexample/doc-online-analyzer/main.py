from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from openai import OpenAI
import requests
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
import io

load_dotenv()

app = FastAPI(title="Dropbox Document Analyzer", description="Analyze documents from Dropbox shared links using OpenAI")

class AnalyzeRequest(BaseModel):
    url: str
    question: str

# OpenAI client with OpenRouter
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=os.getenv("OPENROUTER_API_KEY", "TU_API_KEY_DE_OPENROUTER_AQUI")
)

def download_file_content(url):
    # Convert shared link to direct download link
    from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
    parsed = urlparse(url)
    query = parse_qs(parsed.query)
    query['dl'] = ['1']
    parsed = parsed._replace(query=urlencode(query, doseq=True))
    url = urlunparse(parsed)
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
    response = requests.get(url, headers=headers, allow_redirects=True)
    print(f"Status: {response.status_code}, URL: {response.url}, Content length: {len(response.content)}")
    if response.status_code != 200:
        print(f"Response text: {response.text}")
        raise HTTPException(status_code=400, detail="Failed to download file from Dropbox")
    return response.content

import zipfile

def extract_text(content, filename):
    # Check if it's a zip file (e.g., folder download)
    if zipfile.is_zipfile(io.BytesIO(content)):
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            print(f"Zip contents: {zf.namelist()}")
            text = ''
            for name in zf.namelist():
                ext = name.split('.')[-1].lower()
                if ext in ['txt', 'pdf', 'docx', 'csv']:
                    with zf.open(name) as f:
                        content_file = f.read()
                        if ext in ['txt', 'csv']:
                            try:
                                text += content_file.decode('utf-8') + '\n'
                            except UnicodeDecodeError:
                                try:
                                    text += content_file.decode('latin-1') + '\n'
                                except UnicodeDecodeError:
                                    text += content_file.decode('cp1252', errors='ignore') + '\n'
                        elif ext == 'pdf':
                            pdf = PdfReader(io.BytesIO(content_file))
                            for page in pdf.pages:
                                text += page.extract_text() + '\n'
                        elif ext == 'docx':
                            doc = Document(io.BytesIO(content_file))
                            for para in doc.paragraphs:
                                text += para.text + '\n'
            if text:
                return text
            else:
                raise HTTPException(status_code=400, detail="No supported files (TXT, PDF, DOCX, CSV) found in the zip")
    
    # Otherwise, treat as single file
    ext = filename.split('.')[-1].lower()
    if ext == 'pdf':
        pdf = PdfReader(io.BytesIO(content))
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
        return text
    elif ext == 'docx':
        doc = Document(io.BytesIO(content))
        text = ''
        for para in doc.paragraphs:
            text += para.text + '\n'
        return text
    elif ext in ['txt', 'csv']:
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except UnicodeDecodeError:
                return content.decode('cp1252', errors='ignore')
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Supported: PDF, DOCX, TXT, CSV, or ZIP with supported files")

@app.post("/analyze")
async def analyze_document(request: AnalyzeRequest):
    try:
        # Extract filename from URL
        filename = request.url.split('/')[-1].split('?')[0]
        content = download_file_content(request.url)
        text = extract_text(content, filename)
        text = text[:4000]  # Limit for API

        # Analyze with AI
        response = client.chat.completions.create(
            model="meta-llama/llama-3.2-3b-instruct:free",
            messages=[
                {"role": "user", "content": f"Responde a la pregunta: '{request.question}' basada en el contenido del documento: {text}. Responde en español."}
            ],
            max_tokens=200
        )
        answer = response.choices[0].message.content
        return {"answer": answer}
    except Exception as e:
        print(f"Error: {type(e)} {e}")
        if hasattr(e, 'detail'):
            print(f"Detail: {e.detail}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def root():
    return {"message": "Dropbox Document Analyzer is running. Use POST /analyze with {'url': 'dropbox_shared_link', 'question': 'your question'}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8003)